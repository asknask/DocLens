"""
DOCX document extraction using python-docx.
Extracts paragraphs, tables, and embedded images.
"""
import io
import hashlib
from datetime import datetime
from pathlib import Path
from typing import Optional
from zipfile import ZipFile

from docx import Document as DocxDocument
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from PIL import Image

from app.models.ir_models import (
    DocumentIR,
    DocumentMetadata,
    TextBlock,
    TableBlock,
    ImageBlock,
    BlockLocation,
    TableCell,
)


def extract_docx(
    file_bytes: bytes,
    filename: str,
    job_dir: Optional[Path] = None,
    extract_images: bool = True,
) -> DocumentIR:
    """
    Extract content from a DOCX file.
    
    Args:
        file_bytes: DOCX file content
        filename: Original filename
        job_dir: Optional directory to save extracted images
        extract_images: Whether to extract embedded images
        
    Returns:
        DocumentIR with extracted content
    """
    doc = DocxDocument(io.BytesIO(file_bytes))
    
    blocks = []
    total_chars = 0
    image_counter = 0
    
    # Track paragraph indices for block ordering
    paragraph_idx = 0
    
    # Process paragraphs and tables in document order
    for element in doc.element.body:
        tag = element.tag.split('}')[-1]  # Get tag name without namespace
        
        if tag == 'p':  # Paragraph
            # Find corresponding paragraph in doc.paragraphs
            if paragraph_idx < len(doc.paragraphs):
                para = doc.paragraphs[paragraph_idx]
                text = para.text.strip()
                
                if text:
                    char_count = len(text)
                    total_chars += char_count
                    
                    # Extract style info
                    style = {}
                    if para.style:
                        style["style_name"] = para.style.name
                    
                    blocks.append(TextBlock(
                        content=text,
                        location=BlockLocation(paragraph_index=paragraph_idx),
                        style=style,
                        char_count=char_count,
                    ))
                
                paragraph_idx += 1
        
        elif tag == 'tbl':  # Table
            # Find corresponding table
            table_idx = len([b for b in blocks if isinstance(b, TableBlock)])
            if table_idx < len(doc.tables):
                table = doc.tables[table_idx]
                table_block = _extract_table(table, paragraph_idx)
                if table_block:
                    blocks.append(table_block)
                    # Add table text to total chars
                    for row in table_block.rows:
                        for cell in row:
                            total_chars += len(cell.content)
    
    # Extract embedded images
    image_blocks = []
    if extract_images and job_dir:
        image_blocks = _extract_images(file_bytes, job_dir)
        blocks.extend(image_blocks)
    
    # Build metadata
    core_props = {}
    if doc.core_properties:
        try:
            core_props = {
                "title": doc.core_properties.title or None,
                "author": doc.core_properties.author or None,
                "subject": doc.core_properties.subject or None,
                "created": doc.core_properties.created.isoformat() if doc.core_properties.created else None,
                "modified": doc.core_properties.modified.isoformat() if doc.core_properties.modified else None,
            }
            # Remove None values
            core_props = {k: v for k, v in core_props.items() if v is not None}
        except Exception:
            pass
    
    metadata = DocumentMetadata(
        filename=filename,
        file_type="docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        size_bytes=len(file_bytes),
        page_count=1,  # DOCX doesn't have page concept without rendering
        total_chars=total_chars,
        total_images=len(image_blocks),
        docx_core_properties=core_props,
    )
    
    # Concatenate all text for full_text
    full_text = "\n\n".join(
        block.content for block in blocks 
        if isinstance(block, TextBlock)
    )
    
    return DocumentIR(
        metadata=metadata,
        blocks=blocks,
        full_text=full_text,
    )


def _extract_table(table, paragraph_idx: int) -> Optional[TableBlock]:
    """Extract a table from a python-docx table object."""
    try:
        rows = []
        for row_idx, row in enumerate(table.rows):
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                cells.append(TableCell(
                    content=cell_text,
                    is_header=(row_idx == 0),
                ))
            rows.append(cells)
        
        if rows:
            return TableBlock(
                rows=rows,
                location=BlockLocation(paragraph_index=paragraph_idx),
            )
    except Exception:
        pass
    
    return None


def _extract_images(file_bytes: bytes, job_dir: Path) -> list[ImageBlock]:
    """
    Extract embedded images from DOCX file.
    
    Args:
        file_bytes: DOCX file content
        job_dir: Directory to save extracted images
        
    Returns:
        List of ImageBlocks
    """
    image_blocks = []
    seen_hashes = set()
    
    try:
        with ZipFile(io.BytesIO(file_bytes)) as zf:
            # Find image files in the archive
            image_files = [
                name for name in zf.namelist()
                if name.startswith("word/media/") and 
                any(name.lower().endswith(ext) for ext in [".png", ".jpg", ".jpeg", ".gif", ".webp"])
            ]
            
            for idx, image_name in enumerate(image_files):
                try:
                    image_data = zf.read(image_name)
                    
                    # Calculate hash for deduplication
                    image_hash = hashlib.md5(image_data).hexdigest()
                    if image_hash in seen_hashes:
                        continue
                    seen_hashes.add(image_hash)
                    
                    # Determine MIME type
                    ext = Path(image_name).suffix.lower()
                    mime_map = {
                        ".png": "image/png",
                        ".jpg": "image/jpeg",
                        ".jpeg": "image/jpeg",
                        ".gif": "image/gif",
                        ".webp": "image/webp",
                    }
                    mime_type = mime_map.get(ext, "image/jpeg")
                    
                    # Get image dimensions
                    width, height = None, None
                    try:
                        img = Image.open(io.BytesIO(image_data))
                        width, height = img.size
                    except Exception:
                        pass
                    
                    # Save image
                    image_filename = f"image_{idx + 1:03d}{ext}"
                    image_path = job_dir / image_filename
                    image_path.write_bytes(image_data)
                    
                    image_blocks.append(ImageBlock(
                        image_id=f"docx_image_{idx + 1}",
                        image_path=str(image_path),
                        image_data=image_data,
                        mime_type=mime_type,
                        width=width,
                        height=height,
                        image_hash=image_hash,
                    ))
                    
                except Exception:
                    continue
    
    except Exception:
        pass
    
    return image_blocks
