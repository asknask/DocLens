"""
File validation for uploaded documents.
Enforces size limits, page counts, and content restrictions.
"""
from dataclasses import dataclass
from typing import Any

import fitz  # PyMuPDF
from docx import Document as DocxDocument
import io

from app.config import get_settings
from .mime import detect_mime_type, get_file_category, is_allowed_mime_type, validate_docx_mime


class ValidationError(Exception):
    """Raised when file validation fails."""
    
    def __init__(self, code: str, message: str, details: dict[str, Any] | None = None):
        self.code = code
        self.message = message
        self.details = details or {}
        super().__init__(message)


@dataclass
class ValidationResult:
    """Result of file validation."""
    valid: bool
    mime_type: str
    file_type: str  # pdf, docx, image
    size_bytes: int
    page_count: int | None = None
    char_count: int | None = None
    image_count: int | None = None
    error_code: str | None = None
    error_message: str | None = None


def validate_file(file_bytes: bytes, filename: str) -> ValidationResult:
    """
    Validate an uploaded file for processing.
    
    Checks:
    - MIME type is allowed
    - File size within limits
    - Page count within limits (PDF)
    - Character count within limits (DOCX)
    - Image count within limits (DOCX)
    
    Args:
        file_bytes: Complete file content
        filename: Original filename
        
    Returns:
        ValidationResult with file metadata
        
    Raises:
        ValidationError: If validation fails
    """
    settings = get_settings()
    size_bytes = len(file_bytes)
    
    # Detect MIME type
    mime_type = detect_mime_type(file_bytes)
    if not mime_type:
        raise ValidationError(
            code="unsupported_file_type",
            message=f"Could not detect file type for '{filename}'",
            details={"filename": filename}
        )
    
    if not is_allowed_mime_type(mime_type):
        raise ValidationError(
            code="unsupported_file_type",
            message=f"File type '{mime_type}' is not supported",
            details={"mime_type": mime_type, "filename": filename}
        )
    
    file_type = get_file_category(mime_type)
    if not file_type:
        raise ValidationError(
            code="unsupported_file_type",
            message=f"Unknown file category for '{mime_type}'",
            details={"mime_type": mime_type}
        )
    
    # Check file size
    max_size = settings.get_max_file_size(file_type)
    if size_bytes > max_size:
        raise ValidationError(
            code="file_too_large",
            message=f"File size {size_bytes / (1024*1024):.1f}MB exceeds limit of {max_size / (1024*1024):.1f}MB",
            details={
                "size_bytes": size_bytes,
                "max_bytes": max_size,
                "file_type": file_type
            }
        )
    
    # Type-specific validation
    page_count = None
    char_count = None
    image_count = None
    
    if file_type == "pdf":
        page_count, char_count = _validate_pdf(file_bytes, settings)
    elif file_type == "docx":
        # Additional DOCX validation
        if not validate_docx_mime(file_bytes):
            raise ValidationError(
                code="invalid_docx",
                message="File appears to be a ZIP but not a valid DOCX document",
                details={"filename": filename}
            )
        char_count, image_count = _validate_docx(file_bytes, settings)
    elif file_type == "image":
        page_count = 1  # Images are single page
    elif file_type == "text":
        page_count = 1  # Text files are considered single-page
        char_count = len(file_bytes.decode('utf-8', errors='replace'))

    
    return ValidationResult(
        valid=True,
        mime_type=mime_type,
        file_type=file_type,
        size_bytes=size_bytes,
        page_count=page_count,
        char_count=char_count,
        image_count=image_count
    )


def _validate_pdf(file_bytes: bytes, settings) -> tuple[int, int]:
    """
    Validate PDF file and extract metadata.
    
    Returns:
        Tuple of (page_count, char_count)
        
    Raises:
        ValidationError: If PDF is invalid or exceeds limits
    """
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as e:
        raise ValidationError(
            code="invalid_pdf",
            message=f"Could not parse PDF file: {str(e)}",
            details={"error": str(e)}
        )
    
    try:
        page_count = len(doc)
        
        # Check page limit
        if page_count > settings.max_pdf_pages:
            raise ValidationError(
                code="too_many_pages",
                message=f"PDF has {page_count} pages, exceeds limit of {settings.max_pdf_pages}",
                details={
                    "page_count": page_count,
                    "max_pages": settings.max_pdf_pages
                }
            )
        
        # Count total characters (approximate)
        total_chars = 0
        for page in doc:
            text = page.get_text()
            total_chars += len(text)
        
        return page_count, total_chars
    finally:
        doc.close()


def _validate_docx(file_bytes: bytes, settings) -> tuple[int, int]:
    """
    Validate DOCX file and extract metadata.
    
    Returns:
        Tuple of (char_count, image_count)
        
    Raises:
        ValidationError: If DOCX is invalid or exceeds limits
    """
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
    except Exception as e:
        raise ValidationError(
            code="invalid_docx",
            message=f"Could not parse DOCX file: {str(e)}",
            details={"error": str(e)}
        )
    
    # Count characters
    total_chars = 0
    for para in doc.paragraphs:
        total_chars += len(para.text)
    
    # Also count table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total_chars += len(cell.text)
    
    if total_chars > settings.max_docx_chars:
        raise ValidationError(
            code="too_many_characters",
            message=f"Document has {total_chars:,} characters, exceeds limit of {settings.max_docx_chars:,}",
            details={
                "char_count": total_chars,
                "max_chars": settings.max_docx_chars
            }
        )
    
    # Count images
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_count += 1
    
    if image_count > settings.max_docx_images:
        raise ValidationError(
            code="too_many_images",
            message=f"Document has {image_count} images, exceeds limit of {settings.max_docx_images}",
            details={
                "image_count": image_count,
                "max_images": settings.max_docx_images
            }
        )
    
    return total_chars, image_count
